"""
Complete PDF Editor with Upload Signature & Full Initials Support
All features including signature management with upload and proper initials handling
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, colorchooser, simpledialog
from tkinter.scrolledtext import ScrolledText
import fitz  # PyMuPDF
from PIL import Image, ImageTk, ImageDraw, ImageFont, ImageFilter, ImageEnhance
import io
from typing import Optional, List, Tuple, Dict
import os
from datetime import datetime
import base64
import json
import hashlib
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pymupdf4llm
import re


class PasswordSetupDialog(tk.Toplevel):
    """Dialog to set up a new master password"""

    def __init__(self, parent):
        super().__init__(parent)
        self.password = None
        self.title("Set Master Password")
        self.geometry("400x280")
        self.resizable(False, False)

        # Make modal
        self.transient(parent)
        self.grab_set()

        # Header
        header_frame = ttk.Frame(self, padding="20")
        header_frame.pack(fill=tk.X)

        ttk.Label(header_frame, text="🔒 Secure Your Signatures",
                 font=('Arial', 14, 'bold')).pack()
        ttk.Label(header_frame,
                 text="Set a master password to encrypt your signature library.\n"
                      "This password will be required to access signatures.",
                 justify=tk.CENTER, wraplength=350).pack(pady=10)

        # Password entry
        entry_frame = ttk.Frame(self, padding="20")
        entry_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(entry_frame, text="Master Password:").pack(anchor=tk.W)
        self.password_entry = ttk.Entry(entry_frame, show="*", width=40)
        self.password_entry.pack(fill=tk.X, pady=5)

        ttk.Label(entry_frame, text="Confirm Password:").pack(anchor=tk.W, pady=(10, 0))
        self.confirm_entry = ttk.Entry(entry_frame, show="*", width=40)
        self.confirm_entry.pack(fill=tk.X, pady=5)

        # Show password checkbox
        self.show_password_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(entry_frame, text="Show password",
                       variable=self.show_password_var,
                       command=self.toggle_password).pack(anchor=tk.W, pady=5)

        # Warning
        ttk.Label(entry_frame,
                 text="⚠️ Important: Do not forget this password!\nIt cannot be recovered.",
                 foreground="red", font=('Arial', 9)).pack(pady=10)

        # Buttons
        button_frame = ttk.Frame(self, padding="20")
        button_frame.pack(fill=tk.X)

        ttk.Button(button_frame, text="Cancel", command=self.cancel).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Set Password", command=self.confirm).pack(side=tk.RIGHT)

        # Bind Enter key
        self.password_entry.bind('<Return>', lambda e: self.confirm_entry.focus())
        self.confirm_entry.bind('<Return>', lambda e: self.confirm())

        # Center on parent
        self.password_entry.focus()
        self.wait_window()

    def toggle_password(self):
        """Toggle password visibility"""
        show = "" if self.show_password_var.get() else "*"
        self.password_entry.config(show=show)
        self.confirm_entry.config(show=show)

    def confirm(self):
        """Confirm password setup"""
        password = self.password_entry.get()
        confirm = self.confirm_entry.get()

        if not password:
            messagebox.showwarning("Empty Password", "Please enter a password.")
            return

        if len(password) < 6:
            messagebox.showwarning("Weak Password", "Password must be at least 6 characters long.")
            return

        if password != confirm:
            messagebox.showerror("Mismatch", "Passwords do not match!")
            return

        self.password = password
        self.destroy()

    def cancel(self):
        """Cancel password setup"""
        self.password = None
        self.destroy()


class PasswordUnlockDialog(tk.Toplevel):
    """Dialog to unlock encrypted signatures"""

    def __init__(self, parent):
        super().__init__(parent)
        self.password = None
        self.title("Unlock Signatures")
        self.geometry("400x200")
        self.resizable(False, False)

        # Make modal
        self.transient(parent)
        self.grab_set()

        # Header
        header_frame = ttk.Frame(self, padding="20")
        header_frame.pack(fill=tk.X)

        ttk.Label(header_frame, text="🔒 Signatures Locked",
                 font=('Arial', 14, 'bold')).pack()
        ttk.Label(header_frame,
                 text="Enter your master password to unlock signatures.",
                 justify=tk.CENTER).pack(pady=10)

        # Password entry
        entry_frame = ttk.Frame(self, padding="20")
        entry_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(entry_frame, text="Master Password:").pack(anchor=tk.W)
        self.password_entry = ttk.Entry(entry_frame, show="*", width=40)
        self.password_entry.pack(fill=tk.X, pady=5)

        # Show password checkbox
        self.show_password_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(entry_frame, text="Show password",
                       variable=self.show_password_var,
                       command=self.toggle_password).pack(anchor=tk.W, pady=5)

        # Buttons
        button_frame = ttk.Frame(self, padding="20")
        button_frame.pack(fill=tk.X)

        ttk.Button(button_frame, text="Cancel", command=self.cancel).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="Unlock", command=self.confirm).pack(side=tk.RIGHT)

        # Bind Enter key
        self.password_entry.bind('<Return>', lambda e: self.confirm())

        # Center on parent
        self.password_entry.focus()
        self.wait_window()

    def toggle_password(self):
        """Toggle password visibility"""
        show = "" if self.show_password_var.get() else "*"
        self.password_entry.config(show=show)

    def confirm(self):
        """Confirm and return password"""
        password = self.password_entry.get()

        if not password:
            messagebox.showwarning("Empty Password", "Please enter your password.")
            return

        self.password = password
        self.destroy()

    def cancel(self):
        """Cancel unlock"""
        self.password = None
        self.destroy()


class SignatureStorage:
    """Manages encrypted signature storage and retrieval with master password"""

    def __init__(self, storage_file="signatures.encrypted", parent_window=None):
        self.storage_file = storage_file
        self.signatures: Dict[str, dict] = {}
        self.cipher = None
        self.password_hash = None
        self.salt = None
        self.parent_window = parent_window
        self.is_unlocked = False

        # Try to load or initialize
        self.initialize()

    def initialize(self):
        """Initialize storage - prompt for password"""
        if os.path.exists(self.storage_file):
            # File exists - prompt for password to unlock
            self.prompt_unlock_password()
        else:
            # New file - prompt to set password
            self.prompt_new_password()

    def derive_key(self, password: str, salt: bytes) -> bytes:
        """Derive encryption key from password using PBKDF2"""
        kdf = PBKDF2HMAC(
            algorithm=hashes.SHA256(),
            length=32,
            salt=salt,
            iterations=100000,
        )
        key = base64.urlsafe_b64encode(kdf.derive(password.encode()))
        return key

    def hash_password(self, password: str) -> str:
        """Hash password for verification"""
        return hashlib.sha256(password.encode()).hexdigest()

    def prompt_new_password(self):
        """Prompt user to set a new master password"""
        dialog = PasswordSetupDialog(self.parent_window)

        if dialog.password:
            password = dialog.password
            # Generate salt
            self.salt = os.urandom(16)
            # Derive key
            key = self.derive_key(password, self.salt)
            self.cipher = Fernet(key)
            # Store password hash
            self.password_hash = self.hash_password(password)
            self.is_unlocked = True
            # Save initial empty structure
            self.save_signatures()
        else:
            # User cancelled - create unencrypted fallback
            self.cipher = None
            self.is_unlocked = False

    def prompt_unlock_password(self):
        """Prompt user for password to unlock signatures"""
        # Load salt and hash first
        try:
            with open(self.storage_file, 'rb') as f:
                data = f.read()
                # First 16 bytes: salt
                self.salt = data[:16]
                # Next 64 bytes: password hash (hex string is 64 chars)
                self.password_hash = data[16:80].decode('utf-8')
                # Rest: encrypted data
                self.encrypted_data = data[80:]
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read signature file:\n{str(e)}")
            self.is_unlocked = False
            return

        # Prompt for password
        dialog = PasswordUnlockDialog(self.parent_window)

        if dialog.password:
            password = dialog.password
            # Verify password
            if self.hash_password(password) == self.password_hash:
                # Correct password - derive key
                key = self.derive_key(password, self.salt)
                self.cipher = Fernet(key)
                self.is_unlocked = True
                self.load_signatures()
            else:
                messagebox.showerror("Error", "Incorrect password!\nSignatures remain locked.")
                self.is_unlocked = False
        else:
            # User cancelled
            self.is_unlocked = False

    def load_signatures(self):
        """Load and decrypt signatures from file"""
        if not self.is_unlocked or not self.cipher:
            return

        try:
            with open(self.storage_file, 'rb') as f:
                data = f.read()
                # Skip salt (16 bytes) and hash (64 bytes)
                encrypted_data = data[80:]
                # Decrypt
                decrypted = self.cipher.decrypt(encrypted_data)
                self.signatures = json.loads(decrypted.decode('utf-8'))
        except Exception as e:
            # File doesn't exist or is corrupted - start fresh
            self.signatures = {}

    def save_signatures(self):
        """Encrypt and save signatures to file"""
        if not self.is_unlocked or not self.cipher:
            return

        try:
            # Serialize signatures
            json_data = json.dumps(self.signatures, indent=2)
            # Encrypt
            encrypted = self.cipher.encrypt(json_data.encode('utf-8'))

            # Write: salt + password_hash + encrypted_data
            with open(self.storage_file, 'wb') as f:
                f.write(self.salt)
                f.write(self.password_hash.encode('utf-8'))
                f.write(encrypted)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save signatures:\n{str(e)}")

    def add_signature(self, name: str, signature_data: bytes, sig_type: str = "signature",
                     width: int = 150, height: int = 50):
        """Add a signature to storage"""
        if not self.is_unlocked:
            messagebox.showwarning("Locked", "Signatures are locked. Cannot add signature.")
            return

        b64_data = base64.b64encode(signature_data).decode('utf-8')

        self.signatures[name] = {
            'data': b64_data,
            'type': sig_type,
            'width': width,
            'height': height,
            'created': datetime.now().isoformat()
        }
        self.save_signatures()

    def get_signature(self, name: str) -> Optional[bytes]:
        """Get signature data by name"""
        if not self.is_unlocked:
            return None

        if name in self.signatures:
            b64_data = self.signatures[name]['data']
            return base64.b64decode(b64_data)
        return None

    def delete_signature(self, name: str):
        """Delete a signature"""
        if not self.is_unlocked:
            messagebox.showwarning("Locked", "Signatures are locked. Cannot delete signature.")
            return

        if name in self.signatures:
            del self.signatures[name]
            self.save_signatures()

    def get_all_names(self) -> List[str]:
        """Get all signature names"""
        if not self.is_unlocked:
            return []
        return list(self.signatures.keys())

    def get_signature_info(self, name: str) -> Optional[dict]:
        """Get signature info"""
        if not self.is_unlocked:
            return None
        return self.signatures.get(name)


# [Rest of annotation classes - same as before]
class Annotation:
    """Base class for all annotations"""
    def __init__(self, page_num: int):
        self.page_num = page_num
        self.selected = False

    def draw(self, draw, zoom_level):
        pass

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        return False


class TextAnnotation(Annotation):
    """Text annotation"""
    def __init__(self, page_num: int, x: float, y: float, text: str,
                 fontsize: int = 12, color: Tuple[float, float, float] = (0, 0, 0),
                 fontname: str = "helv"):
        super().__init__(page_num)
        self.x = x
        self.y = y
        self.text = text
        self.fontsize = fontsize
        self.color = color
        self.fontname = fontname

    def draw(self, draw, zoom_level):
        x = self.x * zoom_level
        y = self.y * zoom_level
        size = int(self.fontsize * zoom_level)

        try:
            font = ImageFont.truetype("arial.ttf", size)
        except:
            font = ImageFont.load_default()

        color = tuple(int(c * 255) for c in self.color)
        draw.text((x, y - size), self.text, fill=color, font=font)

        if self.selected:
            bbox = draw.textbbox((x, y - size), self.text, font=font)
            draw.rectangle(bbox, outline="blue", width=2)

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        annot_x = self.x * zoom
        annot_y = self.y * zoom
        text_width = len(self.text) * self.fontsize * zoom * 0.6
        text_height = self.fontsize * zoom * 1.2
        return (annot_x <= x <= annot_x + text_width and
                annot_y - text_height <= y <= annot_y)


class SignatureAnnotation(Annotation):
    """Electronic signature annotation"""
    def __init__(self, page_num: int, x: float, y: float,
                 signature_data: bytes, width: float = 150, height: float = 50):
        super().__init__(page_num)
        self.x = x
        self.y = y
        self.width = width
        self.height = height
        self.signature_data = signature_data
        self.image = None
        self._load_image()

    def _load_image(self):
        """Load signature image from bytes"""
        self.image = Image.open(io.BytesIO(self.signature_data))

    def draw(self, draw, zoom_level):
        x = int(self.x * zoom_level)
        y = int(self.y * zoom_level)
        w = int(self.width * zoom_level)
        h = int(self.height * zoom_level)

        sig_resized = self.image.resize((w, h), Image.Resampling.LANCZOS)
        img = draw._image
        img.paste(sig_resized, (x, y), sig_resized if sig_resized.mode == 'RGBA' else None)

        if self.selected:
            draw.rectangle([x, y, x + w, y + h], outline="blue", width=2)

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        sx = self.x * zoom
        sy = self.y * zoom
        sw = self.width * zoom
        sh = self.height * zoom
        return sx <= x <= sx + sw and sy <= y <= sy + sh


class ShapeAnnotation(Annotation):
    """Shape annotation"""
    def __init__(self, page_num: int, x1: float, y1: float, x2: float, y2: float,
                 shape_type: str = "rectangle", color: Tuple[int, int, int] = (255, 0, 0),
                 thickness: int = 2, fill: bool = False):
        super().__init__(page_num)
        self.x1 = x1
        self.y1 = y1
        self.x2 = x2
        self.y2 = y2
        self.shape_type = shape_type
        self.color = color
        self.thickness = thickness
        self.fill = fill

    def draw(self, draw, zoom_level):
        x1 = int(self.x1 * zoom_level)
        y1 = int(self.y1 * zoom_level)
        x2 = int(self.x2 * zoom_level)
        y2 = int(self.y2 * zoom_level)
        thickness = max(1, int(self.thickness * zoom_level))

        if self.shape_type == "rectangle":
            if self.fill:
                draw.rectangle([x1, y1, x2, y2], fill=self.color, outline=self.color, width=thickness)
            else:
                draw.rectangle([x1, y1, x2, y2], outline=self.color, width=thickness)
        elif self.shape_type == "circle":
            if self.fill:
                draw.ellipse([x1, y1, x2, y2], fill=self.color, outline=self.color, width=thickness)
            else:
                draw.ellipse([x1, y1, x2, y2], outline=self.color, width=thickness)
        elif self.shape_type == "line":
            draw.line([x1, y1, x2, y2], fill=self.color, width=thickness)
        elif self.shape_type == "arrow":
            draw.line([x1, y1, x2, y2], fill=self.color, width=thickness)
            import math
            angle = math.atan2(y2 - y1, x2 - x1)
            arrow_size = 10 * zoom_level
            left_angle = angle + 2.7
            right_angle = angle - 2.7
            left_x = x2 - arrow_size * math.cos(left_angle)
            left_y = y2 - arrow_size * math.sin(left_angle)
            right_x = x2 - arrow_size * math.cos(right_angle)
            right_y = y2 - arrow_size * math.sin(right_angle)
            draw.polygon([x2, y2, left_x, left_y, right_x, right_y], fill=self.color)

        if self.selected:
            draw.rectangle([min(x1, x2) - 2, min(y1, y2) - 2,
                          max(x1, x2) + 2, max(y1, y2) + 2],
                          outline="blue", width=2)

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        sx1, sy1 = self.x1 * zoom, self.y1 * zoom
        sx2, sy2 = self.x2 * zoom, self.y2 * zoom
        margin = 5
        return (min(sx1, sx2) - margin <= x <= max(sx1, sx2) + margin and
                min(sy1, sy2) - margin <= y <= max(sy1, sy2) + margin)


class HighlightAnnotation(Annotation):
    """Highlight annotation"""
    def __init__(self, page_num: int, x1: float, y1: float, x2: float, y2: float,
                 color: Tuple[int, int, int] = (255, 255, 0)):
        super().__init__(page_num)
        self.x1 = x1
        self.y1 = y1
        self.x2 = x2
        self.y2 = y2
        self.color = color

    def draw(self, draw, zoom_level):
        x1 = int(self.x1 * zoom_level)
        y1 = int(self.y1 * zoom_level)
        x2 = int(self.x2 * zoom_level)
        y2 = int(self.y2 * zoom_level)

        overlay = Image.new('RGBA', draw._image.size, (0, 0, 0, 0))
        overlay_draw = ImageDraw.Draw(overlay)
        color_with_alpha = (*self.color, 100)
        overlay_draw.rectangle([x1, y1, x2, y2], fill=color_with_alpha)

        draw._image.paste(Image.alpha_composite(draw._image.convert('RGBA'), overlay).convert('RGB'))

        if self.selected:
            draw.rectangle([x1, y1, x2, y2], outline="blue", width=2)

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        sx1, sy1 = self.x1 * zoom, self.y1 * zoom
        sx2, sy2 = self.x2 * zoom, self.y2 * zoom
        return (min(sx1, sx2) <= x <= max(sx1, sx2) and
                min(sy1, sy2) <= y <= max(sy1, sy2))


class StampAnnotation(Annotation):
    """Stamp annotation"""
    def __init__(self, page_num: int, x: float, y: float, stamp_type: str = "approved"):
        super().__init__(page_num)
        self.x = x
        self.y = y
        self.stamp_type = stamp_type
        self.width = 100
        self.height = 40

    def draw(self, draw, zoom_level):
        x = int(self.x * zoom_level)
        y = int(self.y * zoom_level)
        w = int(self.width * zoom_level)
        h = int(self.height * zoom_level)

        stamps = {
            "approved": ("APPROVED", (0, 150, 0)),
            "rejected": ("REJECTED", (200, 0, 0)),
            "confidential": ("CONFIDENTIAL", (200, 0, 0)),
            "draft": ("DRAFT", (128, 128, 128)),
            "final": ("FINAL", (0, 0, 200)),
            "reviewed": ("REVIEWED", (150, 0, 150)),
        }

        text, color = stamps.get(self.stamp_type, ("APPROVED", (0, 150, 0)))

        draw.rounded_rectangle([x, y, x + w, y + h], radius=5,
                              outline=color, width=3)

        try:
            font_size = int(16 * zoom_level)
            font = ImageFont.truetype("arial.ttf", font_size)
        except:
            font = ImageFont.load_default()

        bbox = draw.textbbox((0, 0), text, font=font)
        text_w = bbox[2] - bbox[0]
        text_h = bbox[3] - bbox[1]
        text_x = x + (w - text_w) // 2
        text_y = y + (h - text_h) // 2

        draw.text((text_x, text_y), text, fill=color, font=font)

        date_text = datetime.now().strftime("%Y-%m-%d")
        try:
            date_font = ImageFont.truetype("arial.ttf", int(10 * zoom_level))
        except:
            date_font = ImageFont.load_default()

        bbox = draw.textbbox((0, 0), date_text, font=date_font)
        date_w = bbox[2] - bbox[0]
        date_x = x + (w - date_w) // 2
        date_y = y + h + 2
        draw.text((date_x, date_y), date_text, fill=color, font=date_font)

        if self.selected:
            draw.rectangle([x - 2, y - 2, x + w + 2, y + h + 15],
                          outline="blue", width=2)

    def contains_point(self, x: float, y: float, zoom: float) -> bool:
        sx = self.x * zoom
        sy = self.y * zoom
        sw = self.width * zoom
        sh = self.height * zoom + 15
        return sx <= x <= sx + sw and sy <= y <= sy + sh


class SignaturePad(tk.Toplevel):
    """Drawing pad for creating signatures"""

    def __init__(self, parent, callback, title="Draw Signature",
                 allow_save=True, storage=None, sig_type="signature"):
        super().__init__(parent)
        self.callback = callback
        self.title(title)
        self.geometry("550x350")
        self.resizable(False, False)
        self.allow_save = allow_save
        self.storage = storage
        self.sig_type = sig_type  # "signature" or "initials"

        # Drawing data
        self.image = Image.new('RGBA', (480, 200), (255, 255, 255, 0))
        self.draw = ImageDraw.Draw(self.image)
        self.last_x = None
        self.last_y = None

        # Canvas
        canvas_frame = ttk.Frame(self, padding="10")
        canvas_frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(canvas_frame, text="Draw your signature below:",
                 font=('Arial', 10)).pack()

        self.canvas = tk.Canvas(canvas_frame, width=480, height=200,
                               bg='white', relief=tk.SUNKEN, bd=2)
        self.canvas.pack(pady=5)

        # Bindings
        self.canvas.bind('<B1-Motion>', self.paint)
        self.canvas.bind('<ButtonRelease-1>', self.reset)

        # Buttons
        button_frame = ttk.Frame(self, padding="10")
        button_frame.pack(fill=tk.X)

        ttk.Button(button_frame, text="Clear", command=self.clear).pack(side=tk.LEFT, padx=5)

        if allow_save and storage:
            ttk.Button(button_frame, text="💾 Save for Later",
                      command=self.save_signature).pack(side=tk.LEFT, padx=5)

        ttk.Button(button_frame, text="Cancel", command=self.cancel).pack(side=tk.RIGHT, padx=5)
        ttk.Button(button_frame, text="✓ Use This",
                  command=self.done).pack(side=tk.RIGHT, padx=5)

    def paint(self, event):
        x, y = event.x, event.y
        if self.last_x and self.last_y:
            self.canvas.create_line(self.last_x, self.last_y, x, y,
                                   width=2, fill='black', capstyle=tk.ROUND, smooth=True)
            self.draw.line([self.last_x, self.last_y, x, y],
                          fill=(0, 0, 0, 255), width=3)
        self.last_x = x
        self.last_y = y

    def reset(self, event):
        self.last_x = None
        self.last_y = None

    def clear(self):
        self.canvas.delete("all")
        self.image = Image.new('RGBA', (480, 200), (255, 255, 255, 0))
        self.draw = ImageDraw.Draw(self.image)

    def save_signature(self):
        """Save signature to storage"""
        name = simpledialog.askstring("Save Signature",
                                     f"Enter a name for this {self.sig_type}:",
                                     parent=self)
        if name:
            buffer = io.BytesIO()
            self.image.save(buffer, format='PNG')
            sig_data = buffer.getvalue()

            # Determine size based on type
            width = 60 if self.sig_type == "initials" else 150
            height = 30 if self.sig_type == "initials" else 50

            self.storage.add_signature(name, sig_data, self.sig_type, width, height)
            messagebox.showinfo("Saved", f"{self.sig_type.title()} '{name}' saved successfully!")

    def done(self):
        buffer = io.BytesIO()
        self.image.save(buffer, format='PNG')
        sig_data = buffer.getvalue()

        if self.callback:
            self.callback(sig_data)
        self.destroy()

    def cancel(self):
        self.destroy()


class SignatureManagerDialog(tk.Toplevel):
    """Dialog to manage saved signatures - ENHANCED with Upload & Initials"""

    def __init__(self, parent, storage: SignatureStorage, callback=None):
        super().__init__(parent)
        self.storage = storage
        self.callback = callback
        self.title("Signature & Initials Manager")
        self.geometry("700x500")

        # Top section with upload button
        top_frame = ttk.Frame(self, padding="10")
        top_frame.pack(fill=tk.X)

        ttk.Label(top_frame, text="Signature & Initials Library",
                 font=('Arial', 12, 'bold')).pack(side=tk.LEFT)

        ttk.Button(top_frame, text="📁 Upload from File",
                  command=self.upload_signature).pack(side=tk.RIGHT, padx=5)

        # Main container
        main_frame = ttk.Frame(self, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Left: List of signatures
        list_frame = ttk.Frame(main_frame)
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        ttk.Label(list_frame, text="Select a signature or initials:").pack(anchor=tk.W)

        # Listbox with scrollbar
        list_scroll_frame = ttk.Frame(list_frame)
        list_scroll_frame.pack(fill=tk.BOTH, expand=True, pady=5)

        scrollbar = ttk.Scrollbar(list_scroll_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.sig_listbox = tk.Listbox(list_scroll_frame, yscrollcommand=scrollbar.set,
                                       font=('Arial', 10))
        self.sig_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.sig_listbox.yview)

        self.sig_listbox.bind('<<ListboxSelect>>', self.on_select)

        # Right: Preview
        preview_frame = ttk.LabelFrame(main_frame, text="Preview", padding="10")
        preview_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=(10, 0))

        self.preview_canvas = tk.Canvas(preview_frame, width=280, height=180,
                                       bg='white', relief=tk.SUNKEN, bd=1)
        self.preview_canvas.pack()

        self.info_label = ttk.Label(preview_frame, text="", justify=tk.LEFT)
        self.info_label.pack(pady=10)

        # Size options when using signature
        size_frame = ttk.LabelFrame(preview_frame, text="Size Options", padding="5")
        size_frame.pack(fill=tk.X, pady=5)

        self.size_var = tk.StringVar(value="original")
        ttk.Radiobutton(size_frame, text="Original Size", variable=self.size_var,
                       value="original").pack(anchor=tk.W)
        ttk.Radiobutton(size_frame, text="Full Signature (150x50)", variable=self.size_var,
                       value="full").pack(anchor=tk.W)
        ttk.Radiobutton(size_frame, text="Initials Size (60x30)", variable=self.size_var,
                       value="initials").pack(anchor=tk.W)

        # Buttons
        button_frame = ttk.Frame(self, padding="10")
        button_frame.pack(fill=tk.X)

        ttk.Button(button_frame, text="✓ Use Selected",
                  command=self.use_signature).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="🗑️ Delete Selected",
                  command=self.delete_signature).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="Close",
                  command=self.destroy).pack(side=tk.RIGHT, padx=5)

        # Load signatures
        self.refresh_list()

    def remove_white_background(self, img):
        """Remove white background from image and make it transparent"""
        # Convert to RGBA if not already
        if img.mode != 'RGBA':
            img = img.convert('RGBA')

        # Get pixel data
        pixels = img.load()
        width, height = img.size

        # Define threshold for "white" (adjustable for near-white colors)
        white_threshold = 240

        # Make white pixels transparent
        for y in range(height):
            for x in range(width):
                r, g, b, a = pixels[x, y]
                # If pixel is white or near-white, make it transparent
                if r >= white_threshold and g >= white_threshold and b >= white_threshold:
                    pixels[x, y] = (r, g, b, 0)  # Set alpha to 0 (transparent)

        return img

    def sharpen_signature(self, img):
        """Sharpen signature image for better clarity"""
        # Convert to RGBA if needed
        if img.mode != 'RGBA':
            img = img.convert('RGBA')

        # Split into channels
        r, g, b, a = img.split()

        # Create RGB composite for sharpening (can't sharpen alpha channel)
        rgb = Image.merge('RGB', (r, g, b))

        # Apply unsharp mask for sharpening
        # radius=2, percent=150, threshold=3
        rgb = rgb.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))

        # Enhance contrast slightly for crisper edges
        enhancer = ImageEnhance.Contrast(rgb)
        rgb = enhancer.enhance(1.2)  # 20% more contrast

        # Merge back with alpha channel
        r, g, b = rgb.split()
        sharpened = Image.merge('RGBA', (r, g, b, a))

        return sharpened

    def upload_signature(self):
        """Upload signature from image file"""
        file_path = filedialog.askopenfilename(
            title="Select Signature Image",
            filetypes=[
                ("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp"),
                ("PNG Files", "*.png"),
                ("JPEG Files", "*.jpg *.jpeg"),
                ("All Files", "*.*")
            ]
        )

        if not file_path:
            return

        try:
            # Load image
            img = Image.open(file_path)

            # Convert to RGBA if needed
            if img.mode != 'RGBA':
                img = img.convert('RGBA')

            # Remove white background (make transparent)
            img = self.remove_white_background(img)

            # Sharpen signature for better clarity
            img = self.sharpen_signature(img)

            # Ask for name
            name = simpledialog.askstring("Name Signature",
                                         "Enter a name for this signature:",
                                         parent=self)
            if not name:
                return

            # Ask for type
            type_dialog = tk.Toplevel(self)
            type_dialog.title("Signature Type")
            type_dialog.geometry("300x150")

            ttk.Label(type_dialog, text="What type is this?",
                     font=('Arial', 10, 'bold')).pack(pady=10)

            type_var = tk.StringVar(value="signature")
            ttk.Radiobutton(type_dialog, text="Full Signature (150x50)",
                           variable=type_var, value="signature").pack(anchor=tk.W, padx=20)
            ttk.Radiobutton(type_dialog, text="Initials (60x30)",
                           variable=type_var, value="initials").pack(anchor=tk.W, padx=20)

            result = {'confirmed': False}

            def confirm():
                result['confirmed'] = True
                type_dialog.destroy()

            ttk.Button(type_dialog, text="OK", command=confirm).pack(pady=10)
            self.wait_window(type_dialog)

            if not result['confirmed']:
                return

            sig_type = type_var.get()

            # Resize to appropriate size
            if sig_type == "initials":
                target_size = (60, 30)
            else:
                target_size = (150, 50)

            # Resize maintaining aspect ratio
            img.thumbnail(target_size, Image.Resampling.LANCZOS)

            # Create new image with target size and paste centered
            final_img = Image.new('RGBA', target_size, (255, 255, 255, 0))
            offset = ((target_size[0] - img.width) // 2, (target_size[1] - img.height) // 2)
            final_img.paste(img, offset, img if img.mode == 'RGBA' else None)

            # Save to storage
            buffer = io.BytesIO()
            final_img.save(buffer, format='PNG')
            sig_data = buffer.getvalue()

            self.storage.add_signature(name, sig_data, sig_type,
                                      target_size[0], target_size[1])

            messagebox.showinfo("Success", f"Uploaded '{name}' successfully!")
            self.refresh_list()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to upload signature:\n{str(e)}")

    def refresh_list(self):
        """Refresh the signature list with type indicators"""
        self.sig_listbox.delete(0, tk.END)
        for name in self.storage.get_all_names():
            info = self.storage.get_signature_info(name)
            sig_type = info.get('type', 'signature') if info else 'signature'

            # Add type indicator
            if sig_type == "initials":
                display_name = f"📝 {name} (Initials)"
            else:
                display_name = f"✍️ {name} (Signature)"

            self.sig_listbox.insert(tk.END, display_name)

    def on_select(self, event):
        """Handle signature selection"""
        selection = self.sig_listbox.curselection()
        if not selection:
            return

        # Get actual name (remove emoji and type indicator)
        display_name = self.sig_listbox.get(selection[0])
        # Extract name from display (remove emoji and text in parentheses)
        if " (Initials)" in display_name:
            name = display_name.replace("📝 ", "").replace(" (Initials)", "")
        else:
            name = display_name.replace("✍️ ", "").replace(" (Signature)", "")

        sig_data = self.storage.get_signature(name)
        info = self.storage.get_signature_info(name)

        if sig_data:
            # Show preview
            self.preview_canvas.delete("all")
            img = Image.open(io.BytesIO(sig_data))
            # Resize to fit
            img.thumbnail((270, 170), Image.Resampling.LANCZOS)
            self.preview_image = ImageTk.PhotoImage(img)
            self.preview_canvas.create_image(140, 90, image=self.preview_image)

            # Show info
            if info:
                created = info.get('created', 'Unknown')
                if 'T' in created:
                    created = created.split('T')[0]
                sig_type = info.get('type', 'signature')
                width = info.get('width', 'Unknown')
                height = info.get('height', 'Unknown')

                info_text = f"Type: {sig_type.title()}\n"
                info_text += f"Size: {width}x{height}\n"
                info_text += f"Created: {created}"
                self.info_label.config(text=info_text)

    def use_signature(self):
        """Use the selected signature"""
        selection = self.sig_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a signature first")
            return

        # Get actual name
        display_name = self.sig_listbox.get(selection[0])
        if " (Initials)" in display_name:
            name = display_name.replace("📝 ", "").replace(" (Initials)", "")
        else:
            name = display_name.replace("✍️ ", "").replace(" (Signature)", "")

        sig_data = self.storage.get_signature(name)
        info = self.storage.get_signature_info(name)

        if sig_data and self.callback:
            # Determine size based on selection
            size_option = self.size_var.get()

            if size_option == "original":
                width = info.get('width', 150)
                height = info.get('height', 50)
            elif size_option == "full":
                width = 150
                height = 50
            else:  # initials
                width = 60
                height = 30

            self.callback(sig_data, width, height)
            self.destroy()

    def delete_signature(self):
        """Delete the selected signature"""
        selection = self.sig_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a signature first")
            return

        # Get actual name
        display_name = self.sig_listbox.get(selection[0])
        if " (Initials)" in display_name:
            name = display_name.replace("📝 ", "").replace(" (Initials)", "")
        else:
            name = display_name.replace("✍️ ", "").replace(" (Signature)", "")

        if messagebox.askyesno("Confirm Delete",
                              f"Are you sure you want to delete '{name}'?"):
            self.storage.delete_signature(name)
            self.refresh_list()
            self.preview_canvas.delete("all")
            self.info_label.config(text="")
            messagebox.showinfo("Deleted", f"'{name}' deleted")


# Import the rest from pdf_editor_interactive for text entry
from pdf_editor_interactive import FloatingTextEntry


class CompletePDFEditor:
    """Complete PDF Editor with all features"""

    def __init__(self, root):
        self.root = root
        self.root.title("Complete PDF Editor Pro - Secure Edition")
        self.root.geometry("1400x900")

        # Initialize signature storage with parent window for password dialogs
        self.signature_storage = SignatureStorage(parent_window=root)

        # PDF state
        self.pdf_document: Optional[fitz.Document] = None
        self.current_page_num: int = 0
        self.total_pages: int = 0
        self.pdf_path: Optional[str] = None
        self.pdf_is_encrypted: bool = False

        # Display state
        self.zoom_level: float = 1.0
        self.current_pixmap = None
        self.current_photo = None
        self.page_rotations: Dict[int, int] = {}  # Track rotation per page (0, 90, 180, 270)

        # Annotations
        self.annotations: List[Annotation] = []
        self.selected_annotation: Optional[Annotation] = None

        # Undo/Redo stacks
        self.undo_stack: List = []
        self.redo_stack: List = []

        # Drawing state
        self.is_drawing: bool = False
        self.is_dragging: bool = False
        self.drag_start_x: int = 0
        self.drag_start_y: int = 0
        self.draw_start_x: int = 0
        self.draw_start_y: int = 0

        # Tool settings
        self.current_tool: str = "select"
        self.text_font: str = "Helvetica"
        self.text_size: int = 12
        self.text_color: Tuple[int, int, int] = (0, 0, 0)
        self.text_color_normalized: Tuple[float, float, float] = (0, 0, 0)
        self.shape_color: Tuple[int, int, int] = (255, 0, 0)
        self.shape_thickness: int = 2
        self.shape_fill: bool = False
        self.highlight_color: Tuple[int, int, int] = (255, 255, 0)
        self.stamp_type: str = "approved"

        # Setup UI
        self.setup_ui()

    # Due to length, I'll continue with setup_ui and other methods...
    # The key changes are in SignatureManagerDialog and the upload/initials support
    # I'll include the critical methods for the main editor class

    def setup_ui(self):
        """Setup complete UI"""
        # Menu bar
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Open PDF (Ctrl+O)", command=self.open_pdf,
                            accelerator="Ctrl+O")
        file_menu.add_command(label="Save (Ctrl+S)", command=self.save_pdf,
                            accelerator="Ctrl+S")
        file_menu.add_command(label="Save As...", command=self.save_pdf_as)
        file_menu.add_separator()
        file_menu.add_command(label="Remove Password Protection", command=self.remove_password_protection)
        file_menu.add_command(label="Bulk Remove Passwords...", command=self.bulk_remove_passwords)
        file_menu.add_command(label="Convert to Word (.docx)", command=self.convert_to_word)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)

        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Undo (Ctrl+Z)", command=self.undo,
                            accelerator="Ctrl+Z")
        edit_menu.add_command(label="Redo (Ctrl+Y)", command=self.redo,
                            accelerator="Ctrl+Y")
        edit_menu.add_separator()
        edit_menu.add_command(label="Delete Selected (Del)", command=self.delete_selected,
                            accelerator="Del")

        sig_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Signatures", menu=sig_menu)
        sig_menu.add_command(label="📁 Upload Signature from File", command=self.upload_signature_direct)
        sig_menu.add_command(label="📝 Manage Signature Library", command=self.manage_signatures)
        sig_menu.add_command(label="✍️ Use Saved Signature", command=self.use_saved_signature)

        view_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="View", menu=view_menu)
        view_menu.add_command(label="Rotate Clockwise (Ctrl+])", command=self.rotate_clockwise,
                             accelerator="Ctrl+]")
        view_menu.add_command(label="Rotate Counter-Clockwise (Ctrl+[)", command=self.rotate_counter_clockwise,
                             accelerator="Ctrl+[")
        view_menu.add_separator()
        view_menu.add_command(label="Reset Rotation", command=self.reset_rotation)

        # Keyboard shortcuts
        self.root.bind('<Control-o>', lambda e: self.open_pdf())
        self.root.bind('<Control-s>', lambda e: self.save_pdf())
        self.root.bind('<Control-z>', lambda e: self.undo())
        self.root.bind('<Control-y>', lambda e: self.redo())
        self.root.bind('<Delete>', lambda e: self.delete_selected())
        self.root.bind('<Control-bracketright>', lambda e: self.rotate_clockwise())
        self.root.bind('<Control-bracketleft>', lambda e: self.rotate_counter_clockwise())

        # Arrow key zoom controls
        self.root.bind('<Up>', lambda e: self.zoom_in())
        self.root.bind('<Down>', lambda e: self.zoom_out())

        # Main container
        main_container = ttk.PanedWindow(self.root, orient=tk.HORIZONTAL)
        main_container.pack(fill=tk.BOTH, expand=True)

        # Left panel - Tools and properties
        left_panel = ttk.Frame(main_container, width=280)
        main_container.add(left_panel, weight=0)

        # Scrollable frame for left panel
        canvas_left = tk.Canvas(left_panel, width=280)
        scrollbar_left = ttk.Scrollbar(left_panel, orient="vertical", command=canvas_left.yview)
        scrollable_frame = ttk.Frame(canvas_left)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas_left.configure(scrollregion=canvas_left.bbox("all"))
        )

        canvas_left.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas_left.configure(yscrollcommand=scrollbar_left.set)

        canvas_left.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar_left.pack(side=tk.RIGHT, fill=tk.Y)

        # Tools
        tools_frame = ttk.LabelFrame(scrollable_frame, text="Tools", padding="10")
        tools_frame.pack(fill=tk.X, padx=5, pady=5)

        self.tool_var = tk.StringVar(value="select")
        tools = [
            ("select", "Select/Move"),
            ("add_text", "Add Text"),
            ("signature", "Draw Signature"),
            ("initials", "Add Initials"),
            ("rectangle", "Rectangle"),
            ("circle", "Circle"),
            ("line", "Line"),
            ("arrow", "Arrow"),
            ("highlight", "Highlight"),
            ("stamp", "Stamp"),
        ]

        for value, label in tools:
            ttk.Radiobutton(tools_frame, text=label, variable=self.tool_var,
                          value=value, command=self.change_tool).pack(anchor=tk.W, pady=2)

        # Text properties
        text_props_frame = ttk.LabelFrame(scrollable_frame, text="Text Properties", padding="10")
        text_props_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(text_props_frame, text="Font:").pack(anchor=tk.W)
        self.font_combo = ttk.Combobox(text_props_frame, values=["Helvetica", "Arial", "Times"],
                                      width=18, state='readonly')
        self.font_combo.set("Helvetica")
        self.font_combo.pack(fill=tk.X, pady=2)
        self.font_combo.bind('<<ComboboxSelected>>', lambda e: setattr(self, 'text_font', self.font_combo.get()))

        ttk.Label(text_props_frame, text="Size:").pack(anchor=tk.W, pady=(10, 0))
        self.size_var = tk.IntVar(value=12)
        size_spinbox = ttk.Spinbox(text_props_frame, from_=6, to=72,
                                  textvariable=self.size_var, width=18)
        size_spinbox.pack(fill=tk.X, pady=2)
        self.size_var.trace('w', lambda *args: setattr(self, 'text_size', self.size_var.get()))

        ttk.Label(text_props_frame, text="Color:").pack(anchor=tk.W, pady=(10, 0))
        color_frame = ttk.Frame(text_props_frame)
        color_frame.pack(fill=tk.X, pady=2)
        self.color_display = tk.Canvas(color_frame, width=30, height=20,
                                      bg="black", relief=tk.SUNKEN, bd=1)
        self.color_display.pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(color_frame, text="Choose", command=self.choose_text_color).pack(side=tk.LEFT)

        # Shape properties
        shape_props_frame = ttk.LabelFrame(scrollable_frame, text="Shape Properties", padding="10")
        shape_props_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(shape_props_frame, text="Color:").pack(anchor=tk.W)
        shape_color_frame = ttk.Frame(shape_props_frame)
        shape_color_frame.pack(fill=tk.X, pady=2)
        self.shape_color_display = tk.Canvas(shape_color_frame, width=30, height=20,
                                            bg="red", relief=tk.SUNKEN, bd=1)
        self.shape_color_display.pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(shape_color_frame, text="Choose", command=self.choose_shape_color).pack(side=tk.LEFT)

        ttk.Label(shape_props_frame, text="Thickness:").pack(anchor=tk.W, pady=(10, 0))
        self.thickness_var = tk.IntVar(value=2)
        ttk.Spinbox(shape_props_frame, from_=1, to=10,
                   textvariable=self.thickness_var, width=18).pack(fill=tk.X, pady=2)
        self.thickness_var.trace('w', lambda *args: setattr(self, 'shape_thickness', self.thickness_var.get()))

        self.fill_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(shape_props_frame, text="Fill shape",
                       variable=self.fill_var).pack(anchor=tk.W, pady=5)
        self.fill_var.trace('w', lambda *args: setattr(self, 'shape_fill', self.fill_var.get()))

        # Highlight properties
        highlight_frame = ttk.LabelFrame(scrollable_frame, text="Highlight Color", padding="10")
        highlight_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Label(highlight_frame, text="Used by Highlight tool",
                 font=('Arial', 8, 'italic')).pack(anchor=tk.W, pady=(0, 5))

        hl_color_frame = ttk.Frame(highlight_frame)
        hl_color_frame.pack(fill=tk.X)
        self.highlight_color_display = tk.Canvas(hl_color_frame, width=30, height=20,
                                                bg="yellow", relief=tk.SUNKEN, bd=1)
        self.highlight_color_display.pack(side=tk.LEFT, padx=(0, 5))
        ttk.Button(hl_color_frame, text="Choose", command=self.choose_highlight_color).pack(side=tk.LEFT)

        # Stamp properties
        stamp_frame = ttk.LabelFrame(scrollable_frame, text="Stamp Type", padding="10")
        stamp_frame.pack(fill=tk.X, padx=5, pady=5)

        self.stamp_var = tk.StringVar(value="approved")
        for stamp in ["approved", "rejected", "confidential", "draft", "final", "reviewed"]:
            ttk.Radiobutton(stamp_frame, text=stamp.title(), variable=self.stamp_var,
                          value=stamp).pack(anchor=tk.W, pady=1)
        self.stamp_var.trace('w', lambda *args: setattr(self, 'stamp_type', self.stamp_var.get()))

        # Signature management - NEW ENHANCED VERSION
        sig_frame = ttk.LabelFrame(scrollable_frame, text="Signature Quick Access", padding="10")
        sig_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(sig_frame, text="📁 Upload from File",
                  command=self.upload_signature_direct).pack(fill=tk.X, pady=2)
        ttk.Button(sig_frame, text="📝 Manage Library",
                  command=self.manage_signatures).pack(fill=tk.X, pady=2)
        ttk.Button(sig_frame, text="✍️ Use Saved",
                  command=self.use_saved_signature).pack(fill=tk.X, pady=2)

        # Show count
        sig_count = len(self.signature_storage.get_all_names())
        self.sig_count_label = ttk.Label(sig_frame,
                                         text=f"{sig_count} saved signature(s)",
                                         font=('Arial', 8, 'italic'))
        self.sig_count_label.pack(pady=5)

        # Navigation
        nav_frame = ttk.LabelFrame(scrollable_frame, text="Navigation", padding="10")
        nav_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(nav_frame, text="◀ Previous Page", command=self.previous_page).pack(fill=tk.X, pady=2)
        self.page_label = ttk.Label(nav_frame, text="0 / 0")
        self.page_label.pack(pady=5)
        ttk.Button(nav_frame, text="Next Page ▶", command=self.next_page).pack(fill=tk.X, pady=2)

        # Zoom
        zoom_frame = ttk.LabelFrame(scrollable_frame, text="Zoom", padding="10")
        zoom_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(zoom_frame, text="Zoom In (↑)", command=self.zoom_in).pack(fill=tk.X, pady=2)
        ttk.Button(zoom_frame, text="Zoom Out (↓)", command=self.zoom_out).pack(fill=tk.X, pady=2)
        ttk.Button(zoom_frame, text="Reset (100%)", command=self.reset_zoom).pack(fill=tk.X, pady=2)
        self.zoom_label = ttk.Label(zoom_frame, text="100%")
        self.zoom_label.pack(pady=5)

        # Rotation
        rotation_frame = ttk.LabelFrame(scrollable_frame, text="Page Rotation", padding="10")
        rotation_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(rotation_frame, text="↻ Rotate Right (Ctrl+])",
                  command=self.rotate_clockwise).pack(fill=tk.X, pady=2)
        ttk.Button(rotation_frame, text="↺ Rotate Left (Ctrl+[)",
                  command=self.rotate_counter_clockwise).pack(fill=tk.X, pady=2)
        ttk.Button(rotation_frame, text="Reset Rotation",
                  command=self.reset_rotation).pack(fill=tk.X, pady=2)
        self.rotation_label = ttk.Label(rotation_frame, text="0°")
        self.rotation_label.pack(pady=5)

        # Security
        security_frame = ttk.LabelFrame(scrollable_frame, text="Security", padding="10")
        security_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(security_frame, text="🔓 Remove Password Protection",
                  command=self.remove_password_protection).pack(fill=tk.X, pady=2)
        ttk.Button(security_frame, text="🔓 Bulk Remove Passwords",
                  command=self.bulk_remove_passwords).pack(fill=tk.X, pady=2)

        # Export
        export_frame = ttk.LabelFrame(scrollable_frame, text="Export", padding="10")
        export_frame.pack(fill=tk.X, padx=5, pady=5)

        ttk.Button(export_frame, text="📄 Convert to Word (.docx)",
                  command=self.convert_to_word).pack(fill=tk.X, pady=2)

        # Center panel - PDF viewer
        center_panel = ttk.Frame(main_container)
        main_container.add(center_panel, weight=1)

        # Canvas
        canvas_frame = ttk.Frame(center_panel)
        canvas_frame.pack(fill=tk.BOTH, expand=True)

        h_scrollbar = ttk.Scrollbar(canvas_frame, orient=tk.HORIZONTAL)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        v_scrollbar = ttk.Scrollbar(canvas_frame, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.canvas = tk.Canvas(canvas_frame, bg='gray75',
                               xscrollcommand=h_scrollbar.set,
                               yscrollcommand=v_scrollbar.set)
        self.canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        h_scrollbar.config(command=self.canvas.xview)
        v_scrollbar.config(command=self.canvas.yview)

        # Canvas bindings
        self.canvas.bind('<Button-1>', self.on_canvas_click)
        self.canvas.bind('<B1-Motion>', self.on_canvas_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_canvas_release)
        self.canvas.bind('<MouseWheel>', self.on_mouse_wheel)

        # Status bar
        self.status_label = ttk.Label(self.root,
                                     text="Ready | Tips: Mouse wheel = navigate pages, Arrow keys = zoom",
                                     relief=tk.SUNKEN, anchor=tk.W)
        self.status_label.pack(side=tk.BOTTOM, fill=tk.X)

    def update_status(self, message: str):
        """Update status bar"""
        self.status_label.config(text=message)
        self.root.update_idletasks()

    def update_signature_count(self):
        """Update signature count display"""
        count = len(self.signature_storage.get_all_names())
        self.sig_count_label.config(text=f"{count} saved signature(s)")

    def change_tool(self):
        """Handle tool change"""
        self.current_tool = self.tool_var.get()
        cursors = {
            "select": "arrow",
            "add_text": "crosshair",
            "signature": "crosshair",
            "initials": "crosshair",
            "rectangle": "crosshair",
            "circle": "crosshair",
            "line": "crosshair",
            "arrow": "crosshair",
            "highlight": "crosshair",
            "stamp": "crosshair",
        }
        self.canvas.config(cursor=cursors.get(self.current_tool, "arrow"))
        self.update_status(f"Tool: {self.current_tool.replace('_', ' ').title()}")

    def choose_text_color(self):
        """Choose text color"""
        color = colorchooser.askcolor(
            color=f"#{self.text_color[0]:02x}{self.text_color[1]:02x}{self.text_color[2]:02x}"
        )
        if color[0]:
            self.text_color = tuple(int(c) for c in color[0])
            self.text_color_normalized = tuple(c / 255.0 for c in self.text_color)
            self.color_display.config(bg=color[1])

    def choose_shape_color(self):
        """Choose shape color"""
        color = colorchooser.askcolor(
            color=f"#{self.shape_color[0]:02x}{self.shape_color[1]:02x}{self.shape_color[2]:02x}"
        )
        if color[0]:
            self.shape_color = tuple(int(c) for c in color[0])
            self.shape_color_display.config(bg=color[1])

    def choose_highlight_color(self):
        """Choose highlight color"""
        color = colorchooser.askcolor(
            color=f"#{self.highlight_color[0]:02x}{self.highlight_color[1]:02x}{self.highlight_color[2]:02x}"
        )
        if color[0]:
            self.highlight_color = tuple(int(c) for c in color[0])
            self.highlight_color_display.config(bg=color[1])

    def remove_white_background(self, img):
        """Remove white background from image and make it transparent"""
        # Convert to RGBA if not already
        if img.mode != 'RGBA':
            img = img.convert('RGBA')

        # Get pixel data
        pixels = img.load()
        width, height = img.size

        # Define threshold for "white" (adjustable for near-white colors)
        white_threshold = 240

        # Make white pixels transparent
        for y in range(height):
            for x in range(width):
                r, g, b, a = pixels[x, y]
                # If pixel is white or near-white, make it transparent
                if r >= white_threshold and g >= white_threshold and b >= white_threshold:
                    pixels[x, y] = (r, g, b, 0)  # Set alpha to 0 (transparent)

        return img

    def sharpen_signature(self, img):
        """Sharpen signature image for better clarity"""
        # Convert to RGBA if needed
        if img.mode != 'RGBA':
            img = img.convert('RGBA')

        # Split into channels
        r, g, b, a = img.split()

        # Create RGB composite for sharpening (can't sharpen alpha channel)
        rgb = Image.merge('RGB', (r, g, b))

        # Apply unsharp mask for sharpening
        # radius=2, percent=150, threshold=3
        rgb = rgb.filter(ImageFilter.UnsharpMask(radius=2, percent=150, threshold=3))

        # Enhance contrast slightly for crisper edges
        enhancer = ImageEnhance.Contrast(rgb)
        rgb = enhancer.enhance(1.2)  # 20% more contrast

        # Merge back with alpha channel
        r, g, b = rgb.split()
        sharpened = Image.merge('RGBA', (r, g, b, a))

        return sharpened

    def upload_signature_direct(self):
        """Direct upload signature from menu"""
        if not self.pdf_document:
            messagebox.showwarning("No PDF", "Please open a PDF first")
            return

        # Open file dialog
        file_path = filedialog.askopenfilename(
            title="Select Signature Image",
            filetypes=[
                ("Image Files", "*.png *.jpg *.jpeg *.gif *.bmp"),
                ("All Files", "*.*")
            ]
        )

        if file_path:
            try:
                img = Image.open(file_path)
                if img.mode != 'RGBA':
                    img = img.convert('RGBA')

                # Remove white background (make transparent)
                img = self.remove_white_background(img)

                # Sharpen signature for better clarity
                img = self.sharpen_signature(img)

                # Ask for type
                type_dialog = tk.Toplevel(self.root)
                type_dialog.title("Signature Type")
                type_dialog.geometry("300x150")

                ttk.Label(type_dialog, text="What type is this?",
                         font=('Arial', 10, 'bold')).pack(pady=10)

                type_var = tk.StringVar(value="signature")
                ttk.Radiobutton(type_dialog, text="Full Signature (150x50)",
                               variable=type_var, value="signature").pack(anchor=tk.W, padx=20)
                ttk.Radiobutton(type_dialog, text="Initials (60x30)",
                               variable=type_var, value="initials").pack(anchor=tk.W, padx=20)

                result = {'confirmed': False}

                def confirm():
                    result['confirmed'] = True
                    type_dialog.destroy()

                ttk.Button(type_dialog, text="OK", command=confirm).pack(pady=10)
                self.root.wait_window(type_dialog)

                if not result['confirmed']:
                    return

                sig_type = type_var.get()

                # Resize to appropriate size
                if sig_type == "initials":
                    target_size = (60, 30)
                else:
                    target_size = (150, 50)

                img.thumbnail(target_size, Image.Resampling.LANCZOS)

                # Create properly sized image
                final_img = Image.new('RGBA', target_size, (255, 255, 255, 0))
                offset = ((target_size[0] - img.width) // 2, (target_size[1] - img.height) // 2)
                final_img.paste(img, offset, img if img.mode == 'RGBA' else None)

                buffer = io.BytesIO()
                final_img.save(buffer, format='PNG')
                sig_data = buffer.getvalue()

                # Store temporarily and switch to signature mode
                self.pending_signature_data = sig_data
                self.pending_sig_width = target_size[0]
                self.pending_sig_height = target_size[1]
                self.update_status(f"Click on PDF to place uploaded {sig_type}")
                self.tool_var.set("signature" if sig_type == "signature" else "initials")
                self.change_tool()

            except Exception as e:
                messagebox.showerror("Error", f"Failed to load image:\n{str(e)}")

    def manage_signatures(self):
        """Open signature manager"""
        def on_signature_selected(sig_data, width, height):
            if sig_data:
                self.pending_signature_data = sig_data
                self.pending_sig_width = width
                self.pending_sig_height = height
                self.update_status(f"Click on PDF to place signature ({width}x{height})")
                self.tool_var.set("signature")
                self.change_tool()

        SignatureManagerDialog(self.root, self.signature_storage, on_signature_selected)
        self.update_signature_count()

    def use_saved_signature(self):
        """Quick use saved signature"""
        if not self.pdf_document:
            messagebox.showwarning("No PDF", "Please open a PDF first")
            return

        if not self.signature_storage.get_all_names():
            messagebox.showinfo("No Signatures",
                              "No saved signatures found.\n\n"
                              "Use 'Upload Signature' or draw one and save it.")
            return

        self.manage_signatures()

    def _ask_pdf_password(self, filename: str) -> Optional[str]:
        """Show a password dialog pre-populated with default ID for encrypted PDFs"""
        dialog = tk.Toplevel(self.root)
        dialog.title("Password Required")
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        result = {"password": None}

        ttk.Label(dialog, text="PDF is password protected.",
                 font=('Arial', 11, 'bold')).pack(pady=(15, 5))
        ttk.Label(dialog, text=f"Enter password to open:\n{filename}",
                 justify=tk.CENTER).pack(pady=(0, 10))

        entry_frame = ttk.Frame(dialog, padding="10")
        entry_frame.pack(fill=tk.X)

        ttk.Label(entry_frame, text="Password:").pack(anchor=tk.W)
        password_entry = ttk.Entry(entry_frame, show="*", width=40)
        password_entry.pack(fill=tk.X, pady=5)
        password_entry.insert(0, "7004295198084")
        password_entry.select_range(0, tk.END)

        def on_ok():
            result["password"] = password_entry.get()
            dialog.destroy()

        def on_cancel():
            dialog.destroy()

        password_entry.bind('<Return>', lambda e: on_ok())
        password_entry.bind('<Escape>', lambda e: on_cancel())

        btn_frame = ttk.Frame(dialog, padding="10")
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Cancel", command=on_cancel).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="OK", command=on_ok).pack(side=tk.RIGHT)

        password_entry.focus()
        dialog.wait_window()
        return result["password"]

    def open_pdf(self):
        """Open PDF file"""
        file_path = filedialog.askopenfilename(
            title="Select PDF File",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )

        if not file_path:
            return

        try:
            if self.pdf_document:
                self.pdf_document.close()

            # Try to open the PDF
            doc = fitz.open(file_path)

            # Check if PDF is encrypted/password protected
            if doc.is_encrypted:
                self.pdf_is_encrypted = True

                # Try to authenticate (some PDFs may have empty password)
                if not doc.authenticate(""):
                    # Need password from user - pre-populate with default ID
                    password = self._ask_pdf_password(os.path.basename(file_path))

                    if password is None:
                        # User cancelled
                        doc.close()
                        return

                    if not doc.authenticate(password):
                        messagebox.showerror("Error", "Incorrect password!\nCould not open PDF.")
                        doc.close()
                        return

                    messagebox.showinfo("Success", "Password accepted! PDF unlocked.")
            else:
                self.pdf_is_encrypted = False

            # PDF opened successfully
            self.pdf_document = doc
            self.pdf_path = file_path
            self.total_pages = len(self.pdf_document)
            self.current_page_num = 0
            self.annotations = []
            self.undo_stack = []
            self.redo_stack = []
            self.zoom_level = 1.0
            self.page_rotations = {}  # Clear rotations for new PDF

            self.display_current_page()

            enc_status = " (Encrypted)" if self.pdf_is_encrypted else ""
            self.update_status(f"Opened: {os.path.basename(file_path)} ({self.total_pages} pages){enc_status}")
            messagebox.showinfo("Success", f"PDF loaded!\n{self.total_pages} pages{enc_status}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to open PDF:\n{str(e)}")

    def display_current_page(self):
        """Display current PDF page"""
        if not self.pdf_document:
            return

        try:
            page = self.pdf_document[self.current_page_num]

            # Get current page rotation
            current_rotation = self.page_rotations.get(self.current_page_num, 0)

            # Apply zoom and rotation
            mat = fitz.Matrix(self.zoom_level, self.zoom_level)

            # Rotate the matrix if needed
            if current_rotation != 0:
                mat = mat.prerotate(current_rotation)

            pix = page.get_pixmap(matrix=mat, alpha=False)

            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))

            # Draw annotations
            img = self.draw_annotations(img)

            self.current_photo = ImageTk.PhotoImage(img)

            self.canvas.delete("all")
            self.canvas.create_image(0, 0, anchor=tk.NW, image=self.current_photo)
            self.canvas.config(scrollregion=(0, 0, pix.width, pix.height))

            self.page_label.config(text=f"{self.current_page_num + 1} / {self.total_pages}")
            self.zoom_label.config(text=f"{int(self.zoom_level * 100)}%")
            self.rotation_label.config(text=f"{current_rotation}°")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to display page:\n{str(e)}")

    def draw_annotations(self, img):
        """Draw all annotations on image"""
        draw = ImageDraw.Draw(img)

        for annot in self.annotations:
            if annot.page_num == self.current_page_num:
                annot.draw(draw, self.zoom_level)

        return img

    def canvas_to_pdf_coords(self, canvas_x: int, canvas_y: int) -> Tuple[float, float]:
        """Convert canvas to PDF coordinates"""
        scroll_x = self.canvas.canvasx(canvas_x)
        scroll_y = self.canvas.canvasy(canvas_y)
        pdf_x = scroll_x / self.zoom_level
        pdf_y = scroll_y / self.zoom_level
        return pdf_x, pdf_y

    def on_canvas_click(self, event):
        """Handle canvas click"""
        if not self.pdf_document:
            return

        pdf_x, pdf_y = self.canvas_to_pdf_coords(event.x, event.y)
        scroll_x = self.canvas.canvasx(event.x)
        scroll_y = self.canvas.canvasy(event.y)

        # Check if using saved/uploaded signature
        if hasattr(self, 'pending_signature_data'):
            sig_data = self.pending_signature_data
            width = getattr(self, 'pending_sig_width', 150)
            height = getattr(self, 'pending_sig_height', 50)
            # Center the signature on cursor position
            centered_x = pdf_x - (width / 2)
            centered_y = pdf_y - (height / 2)
            sig = SignatureAnnotation(self.current_page_num, centered_x, centered_y, sig_data, width, height)
            self.add_annotation(sig)
            del self.pending_signature_data
            if hasattr(self, 'pending_sig_width'):
                del self.pending_sig_width
                del self.pending_sig_height
            return

        if self.current_tool == "add_text":
            self.add_text_annotation(event.x_root, event.y_root, pdf_x, pdf_y)

        elif self.current_tool == "signature":
            self.add_signature(pdf_x, pdf_y)

        elif self.current_tool == "initials":
            self.add_initials(pdf_x, pdf_y)

        elif self.current_tool == "stamp":
            self.add_stamp_annotation(pdf_x, pdf_y)

        elif self.current_tool in ["rectangle", "circle", "line", "arrow"]:
            self.is_drawing = True
            self.draw_start_x = pdf_x
            self.draw_start_y = pdf_y

        elif self.current_tool == "highlight":
            self.is_drawing = True
            self.draw_start_x = pdf_x
            self.draw_start_y = pdf_y

        elif self.current_tool == "select":
            clicked = None
            for annot in reversed(self.annotations):
                if annot.page_num == self.current_page_num:
                    if annot.contains_point(scroll_x, scroll_y, self.zoom_level):
                        clicked = annot
                        break

            if clicked:
                self.selected_annotation = clicked
                self.is_dragging = True
                self.drag_start_x = pdf_x
                self.drag_start_y = pdf_y
                for a in self.annotations:
                    a.selected = False
                clicked.selected = True
            else:
                for a in self.annotations:
                    a.selected = False
                self.selected_annotation = None

            self.display_current_page()

    def on_canvas_drag(self, event):
        """Handle canvas drag"""
        if not self.pdf_document:
            return

        pdf_x, pdf_y = self.canvas_to_pdf_coords(event.x, event.y)

        if self.is_dragging and self.selected_annotation:
            dx = pdf_x - self.drag_start_x
            dy = pdf_y - self.drag_start_y

            if isinstance(self.selected_annotation, (TextAnnotation, SignatureAnnotation, StampAnnotation)):
                self.selected_annotation.x += dx
                self.selected_annotation.y += dy
            elif isinstance(self.selected_annotation, (ShapeAnnotation, HighlightAnnotation)):
                self.selected_annotation.x1 += dx
                self.selected_annotation.y1 += dy
                self.selected_annotation.x2 += dx
                self.selected_annotation.y2 += dy

            self.drag_start_x = pdf_x
            self.drag_start_y = pdf_y
            self.display_current_page()

    def on_canvas_release(self, event):
        """Handle mouse release"""
        if not self.pdf_document:
            return

        pdf_x, pdf_y = self.canvas_to_pdf_coords(event.x, event.y)

        if self.is_drawing and self.current_tool in ["rectangle", "circle", "line", "arrow"]:
            shape = ShapeAnnotation(
                self.current_page_num,
                self.draw_start_x, self.draw_start_y,
                pdf_x, pdf_y,
                self.current_tool,
                self.shape_color,
                self.shape_thickness,
                self.shape_fill
            )
            self.add_annotation(shape)

        elif self.is_drawing and self.current_tool == "highlight":
            highlight = HighlightAnnotation(
                self.current_page_num,
                self.draw_start_x, self.draw_start_y,
                pdf_x, pdf_y,
                self.highlight_color
            )
            self.add_annotation(highlight)

        self.is_drawing = False
        self.is_dragging = False

    def on_mouse_wheel(self, event):
        """Handle mouse wheel for page navigation"""
        # Scroll up = previous page, Scroll down = next page
        if event.delta > 0:
            self.previous_page()
        else:
            self.next_page()

    def add_text_annotation(self, screen_x, screen_y, pdf_x, pdf_y):
        """Add text annotation with floating entry"""
        def callback(text):
            if text and text.strip():
                annot = TextAnnotation(
                    self.current_page_num, pdf_x, pdf_y, text,
                    self.text_size, self.text_color_normalized, self.text_font
                )
                self.add_annotation(annot)

        FloatingTextEntry(self.root, screen_x, screen_y, callback)

    def add_signature(self, pdf_x, pdf_y):
        """Add signature - handle pending uploads"""
        if hasattr(self, 'pending_signature_data'):
            sig_data = self.pending_signature_data
            width = getattr(self, 'pending_sig_width', 150)
            height = getattr(self, 'pending_sig_height', 50)

            # Center the signature on cursor position
            centered_x = pdf_x - (width / 2)
            centered_y = pdf_y - (height / 2)
            sig = SignatureAnnotation(self.current_page_num, centered_x, centered_y,
                                     sig_data, width, height)
            self.add_annotation(sig)

            del self.pending_signature_data
            if hasattr(self, 'pending_sig_width'):
                del self.pending_sig_width
                del self.pending_sig_height
        else:
            def callback(sig_data):
                if sig_data:
                    # Center the signature (default size 150x50)
                    centered_x = pdf_x - 75  # 150/2
                    centered_y = pdf_y - 25  # 50/2
                    sig = SignatureAnnotation(self.current_page_num, centered_x, centered_y, sig_data)
                    self.add_annotation(sig)

            SignaturePad(self.root, callback, "Draw Your Signature",
                        allow_save=True, storage=self.signature_storage, sig_type="signature")

    def add_initials(self, pdf_x, pdf_y):
        """Add initials with proper sizing"""
        def callback(sig_data):
            if sig_data:
                # Center the initials (size 60x30)
                centered_x = pdf_x - 30  # 60/2
                centered_y = pdf_y - 15  # 30/2
                sig = SignatureAnnotation(self.current_page_num, centered_x, centered_y, sig_data, 60, 30)
                self.add_annotation(sig)

        SignaturePad(self.root, callback, "Draw Your Initials",
                    allow_save=True, storage=self.signature_storage, sig_type="initials")

    def add_stamp_annotation(self, pdf_x, pdf_y):
        """Add stamp"""
        stamp = StampAnnotation(self.current_page_num, pdf_x, pdf_y, self.stamp_type)
        self.add_annotation(stamp)

    def add_annotation(self, annotation):
        """Add annotation to list and update display"""
        self.annotations.append(annotation)
        self.undo_stack.append(('add', annotation))
        self.redo_stack.clear()
        self.display_current_page()
        self.update_status(f"Added {type(annotation).__name__}")

    def delete_selected(self):
        """Delete selected annotation"""
        if self.selected_annotation and self.selected_annotation in self.annotations:
            self.annotations.remove(self.selected_annotation)
            self.undo_stack.append(('delete', self.selected_annotation))
            self.redo_stack.clear()
            self.selected_annotation = None
            self.display_current_page()
            self.update_status("Deleted annotation")

    def undo(self):
        """Undo last action"""
        if not self.undo_stack:
            return

        action, annotation = self.undo_stack.pop()
        self.redo_stack.append((action, annotation))

        if action == 'add':
            self.annotations.remove(annotation)
        elif action == 'delete':
            self.annotations.append(annotation)

        self.display_current_page()
        self.update_status("Undo")

    def redo(self):
        """Redo last undone action"""
        if not self.redo_stack:
            return

        action, annotation = self.redo_stack.pop()
        self.undo_stack.append((action, annotation))

        if action == 'add':
            self.annotations.append(annotation)
        elif action == 'delete':
            self.annotations.remove(annotation)

        self.display_current_page()
        self.update_status("Redo")

    def zoom_in(self):
        """Zoom in"""
        self.zoom_level = min(5.0, self.zoom_level + 0.25)
        self.display_current_page()

    def zoom_out(self):
        """Zoom out"""
        self.zoom_level = max(0.25, self.zoom_level - 0.25)
        self.display_current_page()

    def reset_zoom(self):
        """Reset zoom"""
        self.zoom_level = 1.0
        self.display_current_page()

    def previous_page(self):
        """Go to previous page"""
        if self.pdf_document and self.current_page_num > 0:
            self.current_page_num -= 1
            self.display_current_page()

    def next_page(self):
        """Go to next page"""
        if self.pdf_document and self.current_page_num < self.total_pages - 1:
            self.current_page_num += 1
            self.display_current_page()

    def rotate_clockwise(self):
        """Rotate current page 90 degrees clockwise"""
        if not self.pdf_document:
            return

        current_rotation = self.page_rotations.get(self.current_page_num, 0)
        new_rotation = (current_rotation + 90) % 360
        self.page_rotations[self.current_page_num] = new_rotation
        self.display_current_page()
        self.update_status(f"Rotated page {self.current_page_num + 1} to {new_rotation}°")

    def rotate_counter_clockwise(self):
        """Rotate current page 90 degrees counter-clockwise"""
        if not self.pdf_document:
            return

        current_rotation = self.page_rotations.get(self.current_page_num, 0)
        new_rotation = (current_rotation - 90) % 360
        self.page_rotations[self.current_page_num] = new_rotation
        self.display_current_page()
        self.update_status(f"Rotated page {self.current_page_num + 1} to {new_rotation}°")

    def reset_rotation(self):
        """Reset rotation for current page"""
        if not self.pdf_document:
            return

        if self.current_page_num in self.page_rotations:
            del self.page_rotations[self.current_page_num]
        self.display_current_page()
        self.update_status(f"Reset rotation for page {self.current_page_num + 1}")

    def save_pdf(self):
        """Save PDF"""
        if not self.pdf_document:
            return

        try:
            self.apply_annotations()
            self.pdf_document.save(self.pdf_path, incremental=True, encryption=fitz.PDF_ENCRYPT_KEEP)
            self.update_status("Saved")
            messagebox.showinfo("Success", "PDF saved!")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save:\n{str(e)}")

    def save_pdf_as(self):
        """Save PDF as"""
        if not self.pdf_document:
            return

        output_path = filedialog.asksaveasfilename(
            title="Save PDF As",
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf")]
        )

        if output_path:
            try:
                self.apply_annotations()
                self.pdf_document.save(output_path)
                self.pdf_path = output_path
                self.update_status(f"Saved as: {os.path.basename(output_path)}")
                messagebox.showinfo("Success", "PDF saved!")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save:\n{str(e)}")

    def remove_password_protection(self):
        """Remove password protection from PDF and save without encryption"""
        if not self.pdf_document:
            messagebox.showwarning("No PDF", "Please open a PDF first.")
            return

        # Check if PDF is encrypted
        if not self.pdf_is_encrypted:
            messagebox.showinfo("Not Encrypted",
                              "This PDF is not password protected.\n"
                              "No action needed.")
            return

        # Build output path: same folder, same name with _unprotected suffix
        base, ext = os.path.splitext(self.pdf_path)
        output_path = f"{base}_unprotected{ext}"

        try:
            # Apply any pending annotations
            self.apply_annotations()

            # Save without encryption
            self.pdf_document.save(output_path, encryption=fitz.PDF_ENCRYPT_NONE)

            self.update_status(f"Password removed — saved as: {os.path.basename(output_path)}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to remove password protection:\n{str(e)}")

    def bulk_remove_passwords(self):
        """Remove password protection from multiple PDF files at once"""
        file_paths = filedialog.askopenfilenames(
            title="Select Password-Protected PDFs",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )

        if not file_paths:
            return

        # Ask for password once using the same pre-populated dialog
        dialog = tk.Toplevel(self.root)
        dialog.title("Bulk Password Removal")
        dialog.geometry("400x200")
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()

        result = {"password": None}

        ttk.Label(dialog, text=f"Unlock {len(file_paths)} PDF file(s)",
                 font=('Arial', 11, 'bold')).pack(pady=(15, 5))
        ttk.Label(dialog, text="Enter the shared password for all selected files:",
                 justify=tk.CENTER).pack(pady=(0, 10))

        entry_frame = ttk.Frame(dialog, padding="10")
        entry_frame.pack(fill=tk.X)

        ttk.Label(entry_frame, text="Password:").pack(anchor=tk.W)
        password_entry = ttk.Entry(entry_frame, show="*", width=40)
        password_entry.pack(fill=tk.X, pady=5)
        password_entry.insert(0, "7004295198084")
        password_entry.select_range(0, tk.END)

        def on_ok():
            result["password"] = password_entry.get()
            dialog.destroy()

        def on_cancel():
            dialog.destroy()

        password_entry.bind('<Return>', lambda e: on_ok())
        password_entry.bind('<Escape>', lambda e: on_cancel())

        btn_frame = ttk.Frame(dialog, padding="10")
        btn_frame.pack(fill=tk.X)
        ttk.Button(btn_frame, text="Cancel", command=on_cancel).pack(side=tk.RIGHT, padx=5)
        ttk.Button(btn_frame, text="OK", command=on_ok).pack(side=tk.RIGHT)

        password_entry.focus()
        dialog.wait_window()

        password = result["password"]
        if password is None:
            return

        # Process each file
        succeeded = []
        failed = []

        for file_path in file_paths:
            try:
                doc = fitz.open(file_path)

                if not doc.is_encrypted:
                    # Not encrypted - skip
                    failed.append((os.path.basename(file_path), "Not password protected"))
                    doc.close()
                    continue

                # Try empty password first, then provided password
                if not doc.authenticate("") and not doc.authenticate(password):
                    failed.append((os.path.basename(file_path), "Incorrect password"))
                    doc.close()
                    continue

                # Save unprotected version
                base, ext = os.path.splitext(file_path)
                output_path = f"{base}_unprotected{ext}"
                doc.save(output_path, encryption=fitz.PDF_ENCRYPT_NONE)
                doc.close()
                succeeded.append(os.path.basename(output_path))

            except Exception as e:
                failed.append((os.path.basename(file_path), str(e)))

        # Show summary
        if failed:
            msg = f"Processed {len(file_paths)} file(s). {len(succeeded)} succeeded.\n\n"
            msg += "Failed:\n"
            for name, reason in failed:
                msg += f"  - {name}: {reason}\n"
            messagebox.showwarning("Bulk Password Removal", msg)
        else:
            self.update_status(f"Bulk password removal complete — {len(succeeded)} file(s) saved")

    def convert_to_word(self):
        """Convert PDF to Word document (.docx) using pymupdf4llm for layout analysis"""
        if not self.pdf_document:
            messagebox.showwarning("No PDF", "Please open a PDF first.")
            return

        # Ask where to save the Word document
        output_path = filedialog.asksaveasfilename(
            title="Save Word Document As",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")],
            initialfile=os.path.basename(self.pdf_path).replace(".pdf", ".docx") if self.pdf_path else "document.docx"
        )

        if not output_path:
            return

        try:
            self.update_status("Analyzing PDF layout... Please wait.")
            self.root.update_idletasks()

            # Use pymupdf4llm for layout-aware extraction
            # This provides much better text extraction with proper layout analysis
            md_text = pymupdf4llm.to_markdown(
                self.pdf_path,
                page_chunks=True,  # Get per-page chunks
                write_images=False,  # We'll handle images separately
                show_progress=False
            )

            self.update_status("Creating Word document...")
            self.root.update_idletasks()

            # Create a new Word document
            doc = DocxDocument()

            total_pages = len(self.pdf_document)

            # Process each page
            if isinstance(md_text, list):
                # Page chunks mode
                for page_num, page_data in enumerate(md_text):
                    self.update_status(f"Converting page {page_num + 1} of {total_pages}...")
                    self.root.update_idletasks()

                    # Get the markdown content for this page
                    if isinstance(page_data, dict):
                        page_md = page_data.get('text', '')
                    else:
                        page_md = str(page_data)

                    # Convert markdown to Word content
                    self._markdown_to_word(doc, page_md, page_num)

                    # Extract images for this page
                    if page_num < len(self.pdf_document):
                        self._extract_images_from_page(doc, self.pdf_document[page_num])

                    # Add page break (except for last page)
                    if page_num < total_pages - 1:
                        doc.add_page_break()
            else:
                # Single text mode - process entire document
                self._markdown_to_word(doc, str(md_text), 0)

                # Extract all images
                for page_num in range(total_pages):
                    self._extract_images_from_page(doc, self.pdf_document[page_num])

            # Save the document
            doc.save(output_path)

            self.update_status(f"Converted: {os.path.basename(output_path)}")
            messagebox.showinfo("Success",
                              f"PDF converted to Word successfully!\n\n"
                              f"Saved as:\n{os.path.basename(output_path)}\n\n"
                              f"Pages converted: {total_pages}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to convert PDF to Word:\n{str(e)}")
            self.update_status("Conversion failed")

    def _markdown_to_word(self, doc, markdown_text: str, page_num: int):
        """Convert markdown text to Word document content with proper formatting"""

        lines = markdown_text.split('\n')
        in_table = False
        table_rows = []
        in_code_block = False
        code_lines = []

        i = 0
        while i < len(lines):
            line = lines[i]

            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    if code_lines:
                        para = doc.add_paragraph()
                        para.style = 'No Spacing'
                        for code_line in code_lines:
                            run = para.add_run(code_line + '\n')
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Handle tables (markdown format: | cell | cell |)
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_rows = []

                # Skip separator lines (|---|---|)
                if re.match(r'^\s*\|[\s\-:|\+]+\|\s*$', line):
                    i += 1
                    continue

                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                # End of table - create Word table
                self._create_word_table(doc, table_rows)
                in_table = False
                table_rows = []
                # Don't increment i, process current line normally

            # Handle headings
            if line.startswith('######'):
                para = doc.add_paragraph(line[6:].strip())
                para.style = 'Heading 6' if 'Heading 6' in [s.name for s in doc.styles] else 'Heading 3'
                i += 1
                continue
            elif line.startswith('#####'):
                para = doc.add_paragraph(line[5:].strip())
                para.style = 'Heading 5' if 'Heading 5' in [s.name for s in doc.styles] else 'Heading 3'
                i += 1
                continue
            elif line.startswith('####'):
                para = doc.add_paragraph(line[4:].strip())
                para.style = 'Heading 4' if 'Heading 4' in [s.name for s in doc.styles] else 'Heading 3'
                i += 1
                continue
            elif line.startswith('###'):
                para = doc.add_paragraph(line[3:].strip())
                para.style = 'Heading 3'
                i += 1
                continue
            elif line.startswith('##'):
                para = doc.add_paragraph(line[2:].strip())
                para.style = 'Heading 2'
                i += 1
                continue
            elif line.startswith('#'):
                para = doc.add_paragraph(line[1:].strip())
                para.style = 'Heading 1'
                i += 1
                continue

            # Handle bullet points
            bullet_match = re.match(r'^(\s*)[\*\-\+]\s+(.+)$', line)
            if bullet_match:
                text = bullet_match.group(2)
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Handle numbered lists
            num_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if num_match:
                text = num_match.group(3)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)
                i += 1
                continue

            # Handle horizontal rules
            if re.match(r'^[\-\*_]{3,}\s*$', line.strip()):
                para = doc.add_paragraph('_' * 50)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # Skip empty lines
            if not line.strip():
                i += 1
                continue

            # Regular text - add as paragraph
            para = doc.add_paragraph()
            self._add_formatted_text(para, line.strip())
            i += 1

        # Handle any remaining table
        if in_table and table_rows:
            self._create_word_table(doc, table_rows)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text to paragraph with markdown formatting (bold, italic, etc.)"""
        if not text:
            return
        text = text.strip()

        # Pattern to match **bold**, *italic*, ***bold italic***, `code`
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|__(.+?)__|_(.+?)_)'

        last_end = 0
        for match in re.finditer(pattern, text):
            # Add text before this match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            # Determine formatting
            full_match = match.group(0)
            if full_match.startswith('***') or full_match.startswith('___'):
                # Bold italic
                content = match.group(2) or match.group(0)[3:-3]
                run = paragraph.add_run(content)
                run.bold = True
                run.italic = True
            elif full_match.startswith('**') or full_match.startswith('__'):
                # Bold
                content = match.group(3) or match.group(6) or full_match[2:-2]
                run = paragraph.add_run(content)
                run.bold = True
            elif full_match.startswith('*') or full_match.startswith('_'):
                # Italic
                content = match.group(4) or match.group(7) or full_match[1:-1]
                run = paragraph.add_run(content)
                run.italic = True
            elif full_match.startswith('`'):
                # Code
                content = match.group(5) or full_match[1:-1]
                run = paragraph.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    def _create_word_table(self, doc, table_rows: list):
        """Create a Word table from parsed markdown table rows"""
        if not table_rows:
            return

        num_rows = len(table_rows)
        num_cols = max(len(row) for row in table_rows)

        if num_cols == 0:
            return

        # Create table
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill cells
        for row_idx, row_data in enumerate(table_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.cell(row_idx, col_idx)
                    # Clear default paragraph
                    cell.text = ""
                    para = cell.paragraphs[0]
                    self._add_formatted_text(para, str(cell_text) if cell_text else "")

                    # Make first row bold (header)
                    if row_idx == 0:
                        for run in para.runs:
                            run.bold = True

        # Add spacing after table
        doc.add_paragraph()

    def _extract_images_from_page(self, doc, page):
        """Extract images from PDF page and add to Word document"""
        try:
            image_list = page.get_images(full=True)
            for img_info in image_list:
                try:
                    xref = img_info[0]
                    base_image = self.pdf_document.extract_image(xref)
                    image_bytes = base_image["image"]

                    # Open and process image
                    img = Image.open(io.BytesIO(image_bytes))

                    # Skip very small images (likely artifacts)
                    if img.width < 50 or img.height < 50:
                        continue

                    # Convert to RGB if necessary
                    if img.mode in ('RGBA', 'P'):
                        img = img.convert('RGB')

                    # Save to bytes
                    img_buffer = io.BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)

                    # Calculate appropriate width (max 6 inches)
                    max_width = 6
                    img_width_inches = min(img.width / 96, max_width)

                    # Add image to document
                    doc.add_picture(img_buffer, width=Inches(img_width_inches))

                except Exception:
                    pass
        except Exception:
            pass

    def apply_annotations(self):
        """Apply annotations and rotations to PDF"""
        # First, apply rotations to pages
        for page_num, rotation in self.page_rotations.items():
            if 0 <= page_num < len(self.pdf_document):
                page = self.pdf_document[page_num]
                # Get current rotation and add our rotation
                current_rotation = page.rotation
                new_rotation = (current_rotation + rotation) % 360
                page.set_rotation(new_rotation)

        # Then apply annotations
        for annot in self.annotations:
            page = self.pdf_document[annot.page_num]

            if isinstance(annot, TextAnnotation):
                page.insert_text((annot.x, annot.y), annot.text,
                               fontsize=annot.fontsize, color=annot.color)

            elif isinstance(annot, SignatureAnnotation):
                rect = fitz.Rect(annot.x, annot.y,
                               annot.x + annot.width, annot.y + annot.height)
                page.insert_image(rect, stream=annot.signature_data)

            elif isinstance(annot, ShapeAnnotation):
                rect = fitz.Rect(annot.x1, annot.y1, annot.x2, annot.y2)
                color = tuple(c / 255.0 for c in annot.color)

                if annot.shape_type == "rectangle":
                    page.draw_rect(rect, color=color, width=annot.thickness)
                elif annot.shape_type == "circle":
                    page.draw_circle((annot.x1 + annot.x2) / 2, (annot.y1 + annot.y2) / 2,
                                   abs(annot.x2 - annot.x1) / 2, color=color, width=annot.thickness)
                elif annot.shape_type == "line":
                    page.draw_line((annot.x1, annot.y1), (annot.x2, annot.y2),
                                 color=color, width=annot.thickness)

            elif isinstance(annot, HighlightAnnotation):
                rect = fitz.Rect(annot.x1, annot.y1, annot.x2, annot.y2)
                page.add_highlight_annot(rect)

            elif isinstance(annot, StampAnnotation):
                page.insert_text((annot.x, annot.y + 20), annot.stamp_type.upper(),
                               fontsize=16, color=(0, 0.5, 0))


def main():
    root = tk.Tk()
    app = CompletePDFEditor(root)
    root.mainloop()


if __name__ == "__main__":
    main()
